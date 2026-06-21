import pandas as pd
import config
from docx import Document
from datetime import datetime
def formato_miles(numero):
    return f"{numero:,}".replace(",", ".")


def main():

    input_path = f"{config.RAW_FOLDER}/{config.RAW_FILE_NAME}"
    output_path = f"{config.DOCUMENTATION_FOLDER}/{config.PROFILE_FILE_NAME}"

    df = pd.read_csv(input_path, sep=";")

    filas, columnas = df.shape
    nulos = df.isnull().sum()
    total_celdas = filas * columnas
    total_nulos = nulos.sum()
    total_no_nulos = total_celdas - total_nulos
    completitud = total_no_nulos / total_celdas
    porcentaje_nulos = total_nulos / total_celdas
    columnas_con_nulos = (nulos > 0).sum()
    tipos = df.dtypes

    print("===== PERFIL DEL DATASET =====")
    print(f"Filas: {filas}")
    print(f"Columnas: {columnas}")

    doc = Document()

    doc.add_heading(
        "Perfil Inicial del Dataset - Tráfico Aéreo",
        level=1
    )

    doc.add_paragraph(
        f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    doc.add_heading(
        "1. Características Generales",
        level=2
    )
    
    doc.add_paragraph(f"Cantidad de filas: {formato_miles(filas)}")
    doc.add_paragraph(f"Cantidad de columnas: {formato_miles(columnas)}")
    doc.add_paragraph("Separador utilizado: ;")

    doc.add_heading(
        "2. Completitud de Datos", 
        level=2
    )

    doc.add_paragraph(f"Total de celdas evaluadas: {formato_miles(total_celdas)}")
    doc.add_paragraph(f"Total de valores no nulos: {formato_miles(total_no_nulos)}")
    doc.add_paragraph(f"Total de valores nulos: {formato_miles(total_nulos)}")
    doc.add_paragraph(f"Columnas con valores nulos: {formato_miles(columnas_con_nulos)}")
    doc.add_paragraph(f"Porcentaje de completitud: {completitud:.4%}".replace(".", ","))

    doc.add_heading(
        "3. Tipos de Datos",
        level=2
    )

    tabla_tipos = doc.add_table(rows=1, cols=2)
    tabla_tipos.style = "Table Grid"

    encabezado = tabla_tipos.rows[0].cells
    encabezado[0].text = "Columna"
    encabezado[1].text = "Tipo"

    for columna, tipo in tipos.items():

        fila = tabla_tipos.add_row().cells

        fila[0].text = str(columna)
        fila[1].text = str(tipo)

    doc.add_heading(
        "4. Calidad de Datos", 
        level=2
    )

    doc.add_heading(
        "4.1 Resumen Ejecutivo", 
        level=3
    )

    doc.add_paragraph(
        f"La fuente presenta un nivel de completitud de "
        f"{completitud:.4%}. ".replace(".", ",")
        + f"Se identificaron únicamente {formato_miles(total_nulos)} valores nulos "
        + f"sobre un total de {formato_miles(total_celdas)} celdas analizadas, "
        + f"equivalentes al {porcentaje_nulos:.4%} de la información evaluada.".replace(".", ",")
    )

    doc.add_paragraph(
        "Los valores faltantes se concentran principalmente en campos geográficos "
        "secundarios asociados a origen y destino, por lo que no representan un "
        "riesgo significativo para el análisis ni para la construcción del modelo dimensional."
    )

    doc.add_heading(
        "4.2 Valores Nulos por Campo", 
        level=3
    )

    tabla_nulos = doc.add_table(rows=1, cols=2)
    tabla_nulos.style = "Table Grid"

    encabezado = tabla_nulos.rows[0].cells
    encabezado[0].text = "Columna"
    encabezado[1].text = "Nulos"

    for columna, cantidad in nulos.items():

        fila = tabla_nulos.add_row().cells

        fila[0].text = str(columna)
        fila[1].text = str(cantidad)

    diccionario_campos = {
    "Año": "Año del registro.",
    "Mes": "Mes del registro.",
    "Cod_Operador": "Código del operador aéreo.",
    "Operador": "Nombre del operador o aerolínea.",
    "Grupo": "Grupo o clasificación del operador.",
    "ORIG_1": "Código de origen principal.",
    "DEST_1": "Código de destino principal.",
    "ORIG_1_N": "Nombre del origen principal.",
    "DEST_1_N": "Nombre del destino principal.",
    "ORIG_1_PAIS": "País del origen principal.",
    "DEST_1_PAIS": "País del destino principal.",
    "ORIG_2": "Código de origen asociado al par de ciudades.",
    "DEST_2": "Código de destino asociado al par de ciudades.",
    "ORIG_2_N": "Nombre del origen asociado al par de ciudades.",
    "DEST_2_N": "Nombre del destino asociado al par de ciudades.",
    "ORIG_2_PAIS": "País del origen asociado al par de ciudades.",
    "DEST_2_PAIS": "País del destino asociado al par de ciudades.",
    "OPER_2": "Sentido de la operación: llegadas o salidas.",
    "NAC": "Clasificación del tráfico: nacional o internacional.",
    "PAX_LIB": "Campo asociado a pasajeros liberados. Pendiente de validación oficial.",
    "PASAJEROS": "Cantidad de pasajeros transportados.",
    "CAR_LIB": "Campo asociado a carga liberada. Pendiente de validación oficial.",
    "CARGA (Ton)": "Carga transportada medida en toneladas.",
    "CORREO": "Correo transportado por vía aérea.",
    "Distancia": "Distancia asociada a la ruta."
    }

    doc.add_heading(
        "5. Diccionario de Campos", 
        level=2
    )

    tabla_diccionario = doc.add_table(rows=1, cols=2)
    tabla_diccionario.style = "Table Grid"

    encabezado = tabla_diccionario.rows[0].cells
    encabezado[0].text = "Campo"
    encabezado[1].text = "Descripción"

    for columna in df.columns:
        fila = tabla_diccionario.add_row().cells
        fila[0].text = str(columna)
        fila[1].text = diccionario_campos.get(columna, "Descripción pendiente.")

    doc.add_heading(
        "6. Observaciones Iniciales",
        level=2
    )

    doc.add_paragraph(
        "El dataset contiene información mensual de tráfico aéreo "
        "con variables temporales, operacionales, geográficas y empresariales. "
        "La calidad general de los datos es alta y presenta muy pocos valores nulos."
    )

    doc.save(output_path)

    print()
    print("Documento generado correctamente")
    print(output_path)


if __name__ == "__main__":
    main()